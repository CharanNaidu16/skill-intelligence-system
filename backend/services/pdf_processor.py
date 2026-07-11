"""Resume text extraction. Raises ValueError with a user-facing message on
failure instead of silently returning an empty string."""

from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(file_path)
    if suffix == ".docx":
        return _from_docx(file_path)
    return _from_txt(file_path)


def _from_pdf(pdf_path: Path) -> str:
    try:
        reader = PdfReader(str(pdf_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception as exc:
        raise ValueError(f"Could not read PDF file: {exc}") from exc
    if not text:
        raise ValueError(
            "No text could be extracted from this PDF. It may be a scanned "
            "image; please upload a text-based PDF or a TXT file."
        )
    return text


def _from_docx(docx_path: Path) -> str:
    try:
        document = Document(str(docx_path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Resumes frequently use table layouts; pull their text too.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise ValueError(
            f"Could not read Word file: {exc}. If this is an old .doc file, "
            "re-save it as .docx and try again."
        ) from exc
    if not text:
        raise ValueError("No text could be extracted from this Word document.")
    return text


def _from_txt(txt_path: Path) -> str:
    try:
        text = txt_path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception as exc:
        raise ValueError(f"Could not read text file: {exc}") from exc
    if not text:
        raise ValueError("The uploaded file is empty.")
    return text
